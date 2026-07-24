"""Document generation tools — owned by the Resume Builder Agent.

Produces a formatted professional resume as a .docx (python-docx) or .txt file
in the configured output directory. Structured sections are passed as lists of
pipe-delimited strings, which keeps the generated function-calling schema flat
and reliable for the model to fill in.
"""

import re
from typing import List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from ..config import OUTPUT_DIR

ACCENT = RGBColor(0x1F, 0x3A, 0x5F)


def _slugify(value: str) -> str:
    slug = re.sub(r"[^a-zA-Z0-9]+", "_", value).strip("_").lower()
    return slug or "resume"


def _split_entry(entry: str, expected: int) -> List[str]:
    parts = [p.strip() for p in entry.split("|")]
    parts += [""] * (expected - len(parts))
    return parts[:expected]


def _add_heading(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(10)
    para.paragraph_format.space_after = Pt(2)
    run = para.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = ACCENT
    bar = doc.add_paragraph()
    bar.paragraph_format.space_before = Pt(0)
    bar.paragraph_format.space_after = Pt(4)
    bar_run = bar.add_run("_" * 68)
    bar_run.font.size = Pt(6)
    bar_run.font.color.rgb = ACCENT


def _build_docx(path, data: dict) -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Pt(40)
        section.left_margin = section.right_margin = Pt(50)

    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_para.paragraph_format.space_after = Pt(0)
    name_run = name_para.add_run(data["full_name"].upper())
    name_run.bold = True
    name_run.font.size = Pt(20)
    name_run.font.color.rgb = ACCENT

    if data.get("job_title"):
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.paragraph_format.space_after = Pt(2)
        title_run = title_para.add_run(data["job_title"])
        title_run.font.size = Pt(11.5)
        title_run.italic = True

    contact_bits = [b for b in (data.get("email"), data.get("phone"), data.get("location")) if b]
    if contact_bits:
        contact = doc.add_paragraph()
        contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_run = contact.add_run("  •  ".join(contact_bits))
        contact_run.font.size = Pt(9.5)

    if data.get("summary"):
        _add_heading(doc, "Professional Summary")
        doc.add_paragraph(data["summary"])

    if data.get("skills"):
        _add_heading(doc, "Skills")
        skills = [s.strip() for s in data["skills"].split(",") if s.strip()]
        doc.add_paragraph("  •  ".join(skills))

    if data.get("experience"):
        _add_heading(doc, "Professional Experience")
        for entry in data["experience"]:
            role, company, dates, description = _split_entry(entry, 4)
            head = doc.add_paragraph()
            head.paragraph_format.space_after = Pt(0)
            role_run = head.add_run(role)
            role_run.bold = True
            if company:
                head.add_run(f" — {company}")
            if dates:
                date_run = head.add_run(f"   ({dates})")
                date_run.italic = True
                date_run.font.size = Pt(9.5)
            for line in filter(None, (d.strip() for d in description.split(";"))):
                doc.add_paragraph(line, style="List Bullet")

    if data.get("education"):
        _add_heading(doc, "Education")
        for entry in data["education"]:
            degree, institution, year = _split_entry(entry, 3)
            para = doc.add_paragraph()
            para.paragraph_format.space_after = Pt(0)
            degree_run = para.add_run(degree)
            degree_run.bold = True
            if institution:
                para.add_run(f" — {institution}")
            if year:
                year_run = para.add_run(f"   ({year})")
                year_run.italic = True
                year_run.font.size = Pt(9.5)

    if data.get("certifications"):
        _add_heading(doc, "Certifications")
        for cert in (c.strip() for c in data["certifications"].split(",")):
            if cert:
                doc.add_paragraph(cert, style="List Bullet")

    # Arbitrary caller-defined sections, so the template is not tied to any one
    # profession: Publications, Clinical Rotations, Exhibitions, Languages...
    for section in data.get("additional_sections", []):
        title, _, items = section.partition("|")
        if not title.strip():
            continue
        _add_heading(doc, title.strip())
        for item in (i.strip() for i in items.split(";")):
            if item:
                doc.add_paragraph(item, style="List Bullet")

    doc.save(path)


def _build_txt(path, data: dict) -> None:
    lines = [data["full_name"].upper()]
    if data.get("job_title"):
        lines.append(data["job_title"])
    contact = [b for b in (data.get("email"), data.get("phone"), data.get("location")) if b]
    if contact:
        lines.append(" | ".join(contact))
    lines.append("=" * 68)

    def section(title: str, body: List[str]) -> None:
        lines.extend(["", title.upper(), "-" * len(title)])
        lines.extend(body)

    if data.get("summary"):
        section("Professional Summary", [data["summary"]])
    if data.get("skills"):
        section("Skills", [", ".join(s.strip() for s in data["skills"].split(",") if s.strip())])
    if data.get("experience"):
        body = []
        for entry in data["experience"]:
            role, company, dates, description = _split_entry(entry, 4)
            body.append(f"{role} — {company} ({dates})".strip(" —()"))
            body.extend(f"  - {d.strip()}" for d in description.split(";") if d.strip())
        section("Professional Experience", body)
    if data.get("education"):
        body = []
        for entry in data["education"]:
            degree, institution, year = _split_entry(entry, 3)
            body.append(f"{degree} — {institution} ({year})".strip(" —()"))
        section("Education", body)
    if data.get("certifications"):
        section("Certifications", [f"  - {c.strip()}" for c in data["certifications"].split(",") if c.strip()])
    for extra in data.get("additional_sections", []):
        title, _, items = extra.partition("|")
        if title.strip():
            section(title.strip(), [f"  - {i.strip()}" for i in items.split(";") if i.strip()])

    path.write_text("\n".join(lines), encoding="utf-8")


def generate_resume_document(
    full_name: str,
    email: str = "",
    phone: str = "",
    location: str = "",
    job_title: str = "",
    summary: str = "",
    skills: str = "",
    experience: Optional[List[str]] = None,
    education: Optional[List[str]] = None,
    certifications: str = "",
    additional_sections: Optional[List[str]] = None,
    file_name: str = "",
    file_format: str = "docx",
) -> dict:
    """Generate a formatted professional resume document and save it to disk.

    Works for any profession. Beyond the standard sections, any number of
    custom sections can be supplied, so a doctor, lawyer, teacher or designer
    gets the sections their field expects rather than a software template.

    Args:
        full_name: Candidate's full name. Required.
        email: Contact email address.
        phone: Contact phone number.
        location: City and country, e.g. "Berlin, Germany".
        job_title: Target role or headline, e.g. "Senior Backend Engineer".
        summary: A 2-4 sentence professional summary paragraph.
        skills: Comma-separated skills, e.g. "Python, Django, AWS, SQL".
        experience: List of roles, each a pipe-delimited string in the format
            "Role | Company | Dates | achievement one; achievement two".
        education: List of qualifications, each pipe-delimited in the format
            "Degree | Institution | Year".
        certifications: Comma-separated certification names.
        additional_sections: Optional list of extra sections, each a string in
            the format "Section Title | item one; item two; item three". Use
            this for anything the standard sections do not cover, e.g.
            "Publications | Paper A; Paper B" or "Languages | English; Hindi".
        file_name: Optional name for the downloaded file. Only set this when the
            user explicitly asks for a specific filename; otherwise leave it
            empty and the document downloads as "resume.docx".
        file_format: Output format, either "docx" or "txt". Defaults to "docx".

    Returns:
        A dict with status "success" and the absolute path of the generated
        document, or status "error" with a message.
    """
    if not full_name or not full_name.strip():
        return {"status": "error", "message": "full_name is required to generate a resume."}

    file_format = (file_format or "docx").lower().lstrip(".")
    if file_format not in ("docx", "txt"):
        return {"status": "error", "message": f"Unsupported file_format '{file_format}'. Use 'docx' or 'txt'."}

    data = {
        "full_name": full_name.strip(),
        "email": email.strip(),
        "phone": phone.strip(),
        "location": location.strip(),
        "job_title": job_title.strip(),
        "summary": summary.strip(),
        "skills": skills,
        "experience": [e for e in (experience or []) if e and e.strip()],
        "education": [e for e in (education or []) if e and e.strip()],
        "certifications": certifications,
        "additional_sections": [s for s in (additional_sections or []) if s and s.strip()],
    }

    # Stored uniquely per candidate so a past download link never resolves to a
    # different person's file; the browser downloads it under `download_name`.
    path = OUTPUT_DIR / f"{_slugify(full_name)}_resume.{file_format}"

    # The name the user sees when downloading. Defaults to "resume.<fmt>"; a
    # custom file_name is honoured only when the user explicitly asked for one.
    requested = re.sub(r"[^A-Za-z0-9._-]+", "_", (file_name or "").strip()).strip("._")
    if requested:
        if not requested.lower().endswith(f".{file_format}"):
            requested = f"{requested}.{file_format}"
        download_name = requested
    else:
        download_name = f"resume.{file_format}"

    try:
        if file_format == "docx":
            _build_docx(path, data)
        else:
            _build_txt(path, data)
    except Exception as exc:
        return {"status": "error", "message": f"Failed to generate document: {type(exc).__name__}: {exc}"}

    sections = [
        name
        for name, present in (
            ("summary", data["summary"]),
            ("skills", data["skills"]),
            ("experience", data["experience"]),
            ("education", data["education"]),
            ("certifications", data["certifications"]),
        )
        if present
    ]
    sections += [
        s.split("|")[0].strip().lower() for s in data["additional_sections"] if "|" in s
    ]
    return {
        "status": "success",
        "message": f"Resume generated for {data['full_name']}. It will download as '{download_name}'.",
        "file_path": str(path),
        "file_name": path.name,
        "download_name": download_name,
        "file_format": file_format,
        "sections_included": sections,
    }
