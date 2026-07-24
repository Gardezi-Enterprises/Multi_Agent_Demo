"""Plain-text extraction from uploaded documents.

Used by the chat interface: when a file is attached, its text is extracted here
and inlined into the message, so the Resume Analyzer Agent receives ordinary
resume text and the tool assignments in the spec stay exactly as specified.

Formats: .txt/.md and friends natively, .docx via python-docx, .pdf via pypdf
when it is installed. Anything else is treated as a binary attachment — still
usable as an email attachment, just not readable as text.
"""

from pathlib import Path
from typing import Tuple

TEXT_SUFFIXES = {".txt", ".md", ".markdown", ".rst", ".csv", ".json", ".log", ".rtf"}
MAX_CHARS = 20_000


def _from_docx(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text.strip() for cell in row.cells))
    return "\n".join(parts)


def _from_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover - depends on environment
        raise RuntimeError(
            "Reading PDF files requires pypdf. Install it with: pip install pypdf"
        ) from exc

    reader = PdfReader(str(path))
    return "\n".join((page.extract_text() or "") for page in reader.pages)


def extract_text(file_path: str) -> Tuple[str, str]:
    """Extract plain text from a document.

    Returns:
        A (text, error) pair. Exactly one is non-empty: on success `error` is
        empty; on failure `text` is empty and `error` explains why.
    """
    path = Path(file_path)
    if not path.exists():
        return "", f"File not found: {file_path}"

    suffix = path.suffix.lower()
    try:
        if suffix == ".docx":
            text = _from_docx(path)
        elif suffix == ".pdf":
            text = _from_pdf(path)
        elif suffix in TEXT_SUFFIXES or not suffix:
            text = path.read_text(encoding="utf-8", errors="replace")
        elif suffix == ".doc":
            return "", (
                "Legacy .doc files cannot be read. Please save as .docx or .pdf "
                "and upload again."
            )
        else:
            return "", f"'{suffix}' files cannot be read as text."
    except Exception as exc:
        return "", f"Could not read {path.name}: {type(exc).__name__}: {exc}"

    text = text.strip()
    if not text:
        return "", (
            f"{path.name} contains no extractable text. If it is a scanned image "
            "PDF, it needs OCR before it can be analysed."
        )
    if len(text) > MAX_CHARS:
        text = text[:MAX_CHARS] + "\n\n[... truncated ...]"
    return text, ""
